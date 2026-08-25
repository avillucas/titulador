#!/usr/bin/env python3
"""
Script de conversión de catálogo Word (.docx) a JSON.
Extrae la tabla del Catálogo de Formación Profesional 2025.
"""

import docx
import json
import os

def convert_docx_to_json(doc_path: str, output_dir: str):
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"No se encontró el archivo: {doc_path}")

    os.makedirs(output_dir, exist_ok=True)
    doc = docx.Document(doc_path)

    records_flat = []
    trayectos_grouped = {}

    current_sector = ""
    prev_trayecto_info = {}

    # Procesar únicamente las primeras 39 tablas (las siguientes son duplicados exactos en el docx)
    max_tables = min(39, len(doc.tables))
    for t_idx in range(max_tables):
        t = doc.tables[t_idx]
        for row in t.rows:
            cells = [" ".join(c.text.strip().replace('\r', '').replace('\n', ' ').split()) for c in row.cells]
            
            # Omitir filas de encabezado
            if 'Código' in cells[0] or 'Nombre del Trayecto' in cells[1]:
                continue

            # Detectar fila de título de Sector
            unique_cells = set(cells)
            if len(unique_cells) == 1 and list(unique_cells)[0] != '':
                current_sector = list(unique_cells)[0]
                prev_trayecto_info = {}
                continue

            # Omitir filas totalmente vacías
            if not any(cells):
                continue

            code = cells[0]
            trayecto_name = cells[1]

            # Rebase / forward-fill si es fila de continuación de módulo entre saltos de página/tabla
            if not code and not trayecto_name:
                if cells[10] or cells[11]:
                    code = prev_trayecto_info.get("Código", "")
                    trayecto_name = prev_trayecto_info.get("Nombre del Trayecto", "")
                    cells[2] = cells[2] or prev_trayecto_info.get("Tipo de certificación", "")
                    cells[3] = cells[3] or prev_trayecto_info.get("Certificación", "")
                    cells[4] = cells[4] or prev_trayecto_info.get("Res Jurisdiccional", "")
                    cells[5] = cells[5] or prev_trayecto_info.get("Res Validez Nacional", "")
                    cells[6] = cells[6] or prev_trayecto_info.get("Cohortes asociadas a Res de V. N.", "")
                    cells[7] = cells[7] or prev_trayecto_info.get("Hs. Cat. Trayecto", "")
                    cells[8] = cells[8] or prev_trayecto_info.get("Hs. Reloj Trayecto", "")
                    cells[9] = cells[9] or prev_trayecto_info.get("Requisitos de Ingreso/Correlatividad", "")
                    cells[14] = cells[14] or prev_trayecto_info.get("Código de Diseño Curricular", "")
                    cells[15] = cells[15] or (prev_trayecto_info.get("Diseño reemplazados", "") if len(cells) > 15 else "")

            current_trayecto_info = {
                "Código": code,
                "Nombre del Trayecto": trayecto_name,
                "Tipo de certificación": cells[2],
                "Certificación": cells[3],
                "Res Jurisdiccional": cells[4],
                "Res Validez Nacional": cells[5],
                "Cohortes asociadas a Res de V. N.": cells[6],
                "Hs. Cat. Trayecto": cells[7],
                "Hs. Reloj Trayecto": cells[8],
                "Requisitos de Ingreso/Correlatividad": cells[9],
                "Código de Diseño Curricular": cells[14],
                "Diseño reemplazados": cells[15] if len(cells) > 15 else ""
            }

            if code or trayecto_name:
                prev_trayecto_info = current_trayecto_info

            flat_record = {
                "Sector": current_sector,
                "Código": code,
                "Nombre del Trayecto": trayecto_name,
                "Tipo de certificación": cells[2],
                "Certificación": cells[3],
                "Res Jurisdiccional": cells[4],
                "Res Validez Nacional": cells[5],
                "Cohortes asociadas a Res de V. N.": cells[6],
                "Hs. Cat. Trayecto": cells[7],
                "Hs. Reloj Trayecto": cells[8],
                "Requisitos de Ingreso/Correlatividad": cells[9],
                "Código de Planificación": cells[10],
                "Denominación del Módulo": cells[11],
                "Hs. Cat. Módulo": cells[12],
                "Hs. Reloj Módulo": cells[13],
                "Código de Diseño Curricular": cells[14],
                "Diseño reemplazados": cells[15] if len(cells) > 15 else ""
            }
            records_flat.append(flat_record)

            key = (current_sector, code, trayecto_name)
            if key not in trayectos_grouped:
                trayectos_grouped[key] = {
                    "Sector": current_sector,
                    **current_trayecto_info,
                    "Módulos Acreditables": []
                }
            
            module_info = {
                "Código de Planificación": cells[10],
                "Denominación del Módulo": cells[11],
                "Hs. Cat. Módulo": cells[12],
                "Hs. Reloj Módulo": cells[13]
            }
            trayectos_grouped[key]["Módulos Acreditables"].append(module_info)

    records_grouped = list(trayectos_grouped.values())

    json_flat_path = os.path.join(output_dir, 'CATALOGO_2025_FP.json')
    json_grouped_path = os.path.join(output_dir, 'CATALOGO_2025_FP_agrupado.json')

    with open(json_flat_path, 'w', encoding='utf-8') as f:
        json.dump(records_flat, f, ensure_ascii=False, indent=2)

    with open(json_grouped_path, 'w', encoding='utf-8') as f:
        json.dump(records_grouped, f, ensure_ascii=False, indent=2)

    print(f"✅ Archivo plano guardado en: {json_flat_path} ({len(records_flat)} registros)")
    print(f"✅ Archivo agrupado guardado en: {json_grouped_path} ({len(records_grouped)} trayectos)")

if __name__ == '__main__':
    doc_file = './ejemplos/CATALOGO 2025 FP.docx'
    utils_directory = './utils'
    convert_docx_to_json(doc_file, utils_directory)
